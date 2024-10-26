import gspread
from oauth2client.service_account import ServiceAccountCredentials
from dragonmapper.transcriptions import numbered_to_accented
from datetime import date

#pip3 install python-docx
from docx import Document
from docx.shared import Pt
from docx.shared import Mm

#which sheet and from which line should we start exporting
ExportSheetNameList=["Sheet8","Sheet9","Sheet10","Sheet12","Sheet14","Sheet15","Sheet16","Sheet17", "Sheet18", "Sheet19","Sheet20", "Sheet21","Sheet22","Sheet23","Sheet24","Sheet25", "Sheet26","Sheet27","Sheet28"]
ExportFromLine=1181
 
def retrieveDoc():
    scope = ['https://spreadsheets.google.com/feeds','https://www.googleapis.com/auth/drive']
    creds = ServiceAccountCredentials.from_json_keyfile_name('C:\\Users\\ascemama\\Documents\\Private\\Chinese\\chinese-vocabulary-storage-d11d329ddf29.json', scope)
    client = gspread.authorize(creds)
    chineseVocabDoc = client.open('ChineseVocab')
    return chineseVocabDoc

def createNewWordDoc():
    document = Document()
    #set A4
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.header_distance = Mm(1)
    section.footer_distance = Mm(1)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)  #100 characters per line
    return document

#for those words that were not learned well in the last weeks
def addWordsToBeReReviewed(ChineseVocabDoc, newDoc):
    lineChinese=""
    lineTraduction=""
    isLastLine=False
    sheets = ChineseVocabDoc.worksheets()
    for sheet in sheets:
        records_data = sheet.get_all_records()
         
        records_length=len(records_data)
        for idx in range(0,records_length):
            if(records_data[idx]["ToBeReviewed"]=="x"):
                #print(records_data[idx]["Traduction"])
                pinyinLen=len(numbered_to_accented(records_data[idx]["Pinyin"]))
                traductionLen=len(records_data[idx]["Traduction"])

                if(idx==(records_length-1)):
                    isLastLine=True

                if ((len(lineChinese)+pinyinLen > 40) or (len(lineTraduction)+traductionLen > 40)):
                    padding=" "*(90-len(lineTraduction))
                    newDoc.add_paragraph(lineTraduction+padding+lineChinese)
                    lineChinese=numbered_to_accented(records_data[idx]["Pinyin"])
                    lineTraduction=records_data[idx]["Traduction"]
                else:
                    if idx==0:
                        lineChinese=numbered_to_accented(records_data[idx]["Pinyin"])
                        lineTraduction=records_data[idx]["Traduction"]
                    if (isLastLine):
                        lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Pinyin"])
                        lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
                    else:
                        lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Pinyin"])
                        lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
        
                if(isLastLine):
                    padding=" "*(90-len(lineTraduction))
                    newDoc.add_paragraph(lineTraduction+padding+lineChinese)

def addNewWords(ChineseVocabDoc, newDoc):
    lineChinese=""
    lineTraduction=""
    isLastLine=False
    sheet_instance=ChineseVocabDoc.worksheet(ExportSheetNameList[-1])
    #print(sheet_instance)
    records_data = sheet_instance.get_all_records()
    records_length=len(records_data)
    for idx in range(ExportFromLine,records_length):
        pinyinLen=len(numbered_to_accented(records_data[idx]["Pinyin"]))
        traductionLen=len(records_data[idx]["Traduction"])
        #print("add new words "+records_data[idx]["Traduction"])
        if(idx==(records_length-1)):
            isLastLine=True

        if ((len(lineChinese)+pinyinLen > 40) or (len(lineTraduction)+traductionLen > 40)):
            padding=" "*(90-len(lineTraduction))
            newDoc.add_paragraph(lineTraduction+padding+lineChinese)
            lineChinese=numbered_to_accented(records_data[idx]["Pinyin"])
            lineTraduction=records_data[idx]["Traduction"]
        else:
            if idx==0:
                lineChinese=numbered_to_accented(records_data[idx]["Pinyin"])
                lineTraduction=records_data[idx]["Traduction"]
            if (isLastLine):
                lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Pinyin"])
                lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
            else:
                lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Pinyin"])
                lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
        
        if(isLastLine):
            padding=" "*(90-len(lineTraduction))
            newDoc.add_paragraph(lineTraduction+padding+lineChinese)
 
### Main

chineseVocabDoc=retrieveDoc()
newDoc=createNewWordDoc()
addWordsToBeReReviewed(chineseVocabDoc,newDoc)
addNewWords(chineseVocabDoc,newDoc)
newDoc.save('./vocab_'+str(date.today())+'.docx')