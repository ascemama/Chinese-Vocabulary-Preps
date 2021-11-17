import gspread
from oauth2client.service_account import ServiceAccountCredentials
from dragonmapper.transcriptions import numbered_to_accented
from datetime import date

#pip3 install python-docx
from docx import Document
from docx.shared import Pt
from docx.shared import Mm

#which sheet and from which line should we start exporting
ExportSheetNb=1
ExportFromLine=0

def retrieveDocJSON():
    scope = ['https://spreadsheets.google.com/feeds','https://www.googleapis.com/auth/drive']
    creds = ServiceAccountCredentials.from_json_keyfile_name('C:\\Users\\ascemama\\Documents\chinese-vocabulary-storage-d11d329ddf29.json', scope)
    client = gspread.authorize(creds)
    sheet = client.open('ChineseVocab')
    sheet_instance = sheet.get_worksheet(ExportSheetNb)
    records_data = sheet_instance.get_all_records()
    # test update data
    #print(sheet_instance.cell(col=3,row=2))
    #print(numbered_to_accented(records_data[2]["Pinyin"]))
    #row=["a","b","c"]
    #index=3
    #sheet_instance.insert_row(row,index)
    return records_data


 

def createWordDoc(DocJSON):
    document = Document()
    #set A4
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.header_distance = Mm(1)
    section.footer_distance = Mm(1)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)  #100 characters per line
    lineChinese=""
    lineTraduction=""
    isLastLine=False
    DocJSONLen=len(DocJSON)
    for idx in range(ExportFromLine,DocJSONLen):
        pinyinLen=len(numbered_to_accented(DocJSON[idx]["Pinyin"]))
        traductionLen=len(DocJSON[idx]["Traduction"])

        if(idx==(DocJSONLen-1)):
            isLastLine=True

        if ((len(lineChinese)+pinyinLen > 40) or (len(lineTraduction)+traductionLen > 40)):
            padding=" "*(90-len(lineTraduction))
            document.add_paragraph(lineTraduction+padding+lineChinese)
            lineChinese=numbered_to_accented(DocJSON[idx]["Pinyin"])
            lineTraduction=DocJSON[idx]["Traduction"]
        else:
            if idx==0:
                lineChinese=numbered_to_accented(DocJSON[idx]["Pinyin"])
                lineTraduction=DocJSON[idx]["Traduction"]
            if (isLastLine):
                lineChinese=lineChinese+"/"+numbered_to_accented(DocJSON[idx]["Pinyin"])
                lineTraduction=lineTraduction+"/"+DocJSON[idx]["Traduction"]
            else:
                lineChinese=lineChinese+"/"+numbered_to_accented(DocJSON[idx]["Pinyin"])
                lineTraduction=lineTraduction+"/"+DocJSON[idx]["Traduction"]
        
        if(isLastLine):
            padding=" "*(90-len(lineTraduction))
            document.add_paragraph(lineTraduction+padding+lineChinese)
    document.save('./vocab_'+str(date.today())+'.docx')




### Main
doc=retrieveDocJSON()
createWordDoc(doc)