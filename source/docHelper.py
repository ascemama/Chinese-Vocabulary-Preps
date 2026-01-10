from dragonmapper.transcriptions import numbered_to_accented

#pip3 install python-docx
from docx import Document
from docx.shared import Pt
from docx.shared import Mm

def createNewWordDoc():
    document = Document()
    #set A4
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.header_distance = Mm(1)
    section.footer_distance = Mm(1)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)  #100 characters per line
    return document

#for those words that were not learned well in the last weeks
def addWordsToBeReReviewed(ChineseVocabDoc, newDoc):
    lineChinese=""
    lineTraduction=""
    isLastLine=False
    sheets = ChineseVocabDoc.worksheets()
    for sheet in sheets:
        records_data = sheet.get_all_records()
         
        records_length=len(records_data)
        for idx in range(0,records_length):
            if(records_data[idx]["ToBeReviewed"]=="x"):
                #print(records_data[idx]["Traduction"])
                SourceLen=len(numbered_to_accented(records_data[idx]["Source"]))
                traductionLen=len(records_data[idx]["Traduction"])

                if(idx==(records_length-1)):
                    isLastLine=True

                if ((len(lineChinese)+SourceLen > 40) or (len(lineTraduction)+traductionLen > 40)):
                    padding=" "*(90-len(lineTraduction))
                    newDoc.add_paragraph(lineTraduction+padding+lineChinese)
                    lineChinese=numbered_to_accented(records_data[idx]["Source"])
                    lineTraduction=records_data[idx]["Traduction"]
                else:
                    if idx==0:
                        lineChinese=numbered_to_accented(records_data[idx]["Source"])
                        lineTraduction=records_data[idx]["Traduction"]
                    if (isLastLine):
                        lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Source"])
                        lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
                    else:
                        lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Source"])
                        lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
        
                if(isLastLine):
                    padding=" "*(90-len(lineTraduction))
                    newDoc.add_paragraph(lineTraduction+padding+lineChinese)

def addNewWords(ChineseVocabDoc, newDoc,exportSheetNameList,exportFromLine):
    lineChinese=""
    lineTraduction=""
    isLastLine=False
    sheet_instance=ChineseVocabDoc.worksheet(exportSheetNameList[-1])
    #print(sheet_instance)
    records_data = sheet_instance.get_all_records()
    records_length=len(records_data)
    for idx in range(exportFromLine,records_length):
        SourceLen=len(numbered_to_accented(records_data[idx]["Source"]))
        traductionLen=len(records_data[idx]["Traduction"])
        #print("add new words "+records_data[idx]["Traduction"])
        if(idx==(records_length-1)):
            isLastLine=True

        if ((len(lineChinese)+SourceLen > 40) or (len(lineTraduction)+traductionLen > 40)):
            padding=" "*(90-len(lineTraduction))
            newDoc.add_paragraph(lineTraduction+padding+lineChinese)
            lineChinese=numbered_to_accented(records_data[idx]["Source"])
            lineTraduction=records_data[idx]["Traduction"]
        else:
            if idx==0:
                lineChinese=numbered_to_accented(records_data[idx]["Source"])
                lineTraduction=records_data[idx]["Traduction"]
            if (isLastLine):
                lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Source"])
                lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
            else:
                lineChinese=lineChinese+"/"+numbered_to_accented(records_data[idx]["Source"])
                lineTraduction=lineTraduction+"/"+records_data[idx]["Traduction"]
        
        if(isLastLine):
            padding=" "*(90-len(lineTraduction))
            newDoc.add_paragraph(lineTraduction+padding+lineChinese)
 